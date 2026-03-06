from io import BytesIO

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from flask import Blueprint, request, send_file

from ..models import RewriteTask
from ..rewrite_service import new_task_id, start_rewrite_task
from ..utils import error_response, success_response
from ..extensions import db

rewrite_bp = Blueprint("rewrite", __name__)


@rewrite_bp.post("/rewrite/start")
def rewrite_start():
    body = request.get_json(silent=True) or {}
    url = (body.get("url") or "").strip()
    article_id = (body.get("articleId") or "").strip() or None
    if not url:
        return error_response(4001, "url 不能为空")

    task_id = new_task_id()
    task = RewriteTask(
        task_id=task_id,
        url=url,
        article_id=article_id,
        status="processing",
        progress=0,
        status_text="任务排队中...",
        time_remaining=8,
    )
    db.session.add(task)
    db.session.commit()
    start_rewrite_task(task_id)
    return success_response({"taskId": task_id})


@rewrite_bp.get("/rewrite/status")
def rewrite_status():
    task_id = (request.args.get("taskId") or "").strip()
    if not task_id:
        return error_response(4001, "taskId 不能为空")
    task = RewriteTask.query.filter_by(task_id=task_id).first()
    if not task:
        return error_response(4004, "改写任务不存在")
    if task.error_message:
        return error_response(5001, f"改写失败：{task.error_message}")

    data = {
        "taskId": task.task_id,
        "status": task.status,
        "progress": task.progress,
        "statusText": task.status_text,
        "timeRemaining": task.time_remaining,
        "sourceHtml": task.source_html or "",
    }
    if task.status == "completed":
        data["result"] = {
            "rewrittenBodyHtml": task.rewritten_body_html or "",
            "originalTitle": task.original_title or "",
            "suggestedTitles": task.suggested_titles or [],
        }
    return success_response(data)


@rewrite_bp.post("/rewrite/export-docx")
def rewrite_export_docx():
    body = request.get_json(silent=True) or {}
    task_id = (body.get("taskId") or "").strip()
    if not task_id:
        return error_response(4001, "taskId 不能为空")

    task = RewriteTask.query.filter_by(task_id=task_id).first()
    if not task:
        return error_response(4004, "改写任务不存在")
    if task.status != "completed":
        return error_response(4001, "改写任务未完成，无法导出")
    if not task.rewritten_body_html:
        return error_response(4001, "改写内容为空，无法导出")

    doc = Document()
    title = task.original_title or "改写文章"
    doc.add_heading(title, level=1)

    suggested = task.suggested_titles or []
    if suggested:
        doc.add_heading("推荐标题", level=2)
        for item in suggested:
            doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("改写正文", level=2)
    _append_html_to_docx(doc, task.rewritten_body_html)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"rewrite_{task.task_id}.docx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _append_html_to_docx(doc: Document, html: str):
    soup = BeautifulSoup(html or "", "html.parser")
    blocks = soup.find_all(["p", "h1", "h2", "h3", "li", "img"])
    for node in blocks:
        name = node.name.lower()
        if name == "img":
            src = (
                node.get("src")
                or node.get("data-src")
                or node.get("data-original")
                or node.get("original-src")
                or ""
            ).strip()
            if not src:
                continue
            if src.startswith("//"):
                src = f"https:{src}"
            try:
                resp = requests.get(src, timeout=12)
                resp.raise_for_status()
                bio = BytesIO(resp.content)
                doc.add_picture(bio, width=Inches(5.4))
            except Exception:
                doc.add_paragraph(f"[图片] {src}")
            continue

        text = node.get_text(" ", strip=True)
        if not text:
            continue
        if name == "h1":
            doc.add_heading(text, level=1)
        elif name == "h2":
            doc.add_heading(text, level=2)
        elif name == "h3":
            doc.add_heading(text, level=3)
        elif name == "li":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
