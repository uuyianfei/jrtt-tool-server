import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import requests
import json
import os
import threading
from datetime import datetime
import re
import time
import math

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告：未安装python-docx库，DOCX生成功能将不可用。请运行：pip install python-docx")


class ArticlePrompts:
    """专业化的文章提示词系统 - 智能段落控制版（增强质量与标题字数限制）"""

    @staticmethod
    def count_paragraphs(text):
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text.strip()) if p.strip()]
        return len(paragraphs)

    @staticmethod
    def get_optimization_prompt(original_content, original_paragraphs):
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', original_content)
        original_len = len(chinese_chars)
        target_desc = "扩写到约1200字" if original_len < 1200 else "保持相近字数（但通过手法大幅降低相似度）"

        if original_paragraphs < 7:
            para_target = f"原文共有 **{original_paragraphs} 个自然段**，文章较短。你的改写结果必须达到 **至少 7 个自然段**（建议控制在 7-9 段），让文章更丰富、更有层次。"
        else:
            para_target = f"原文共有 **{original_paragraphs} 个自然段**，你的改写结果必须保持段落数相近，允许误差 **±2 段**，即最终段落数应在 {original_paragraphs-2} 到 {original_paragraphs+2} 之间。"

        return f"""你是一个干了十年的爆文编辑，最擅长把平平无奇的文章改成朋友圈疯传的爆款。现在交给你一篇稿子，你要用你的老辣经验，把它改得让读者一看就停不下来，读完还想骂两句或者赞两句，忍不住在评论区打字。

【原文】
{original_content[:3000]}{'...（原文太长，只能给你看前三千字）' if len(original_content) > 3000 else ''}

【硬性要求】
1.  **字数**：原文大概{original_len}个中文字。你给我{target_desc}。别差太多。
2.  **段落数**：{para_target}
3.  **相似度**：必须和原文长得完全不一样。用乱序、插叙、倒叙、换人称、加例子、删废话……你可以通过同义词替换、句式变换、增加细节等方式来降低相似度，随便你怎么搞，总之和原文的相似度要低于10%，最好让原作者都认不出来。
4.  **语言风格**：用大白话！口语！就像你跟朋友吹牛聊天那样写。别用那些“与此同时”、“然而”、“此外”、“综上所述”之类的词，太假了，读者一眼就看出是AI写的。句子短一点，别绕。
5.  **爆文要素**：
    *   **标题**：起两三个标题，必须控制在 **24-30个中文字符（包括标点）** 之间。标题要有争议性，能吸引读者点击，但不要太过于哗众取宠。避免使用“震惊”、“真相”等低俗词汇。
    *   **开头**：第一句话就要抓住人。可以用一个反问、一个颠覆常识的观点、或者一个很有画面感的场景。
    *   **正文**：要有冲突，有转折，有能戳中读者痛点或爽点的地方。可以把原文里的一个观点放大，制造一点对立面，让读者站队。
    *   **结尾**：别就那么完了。丢一个问题给读者，引导他们评论。比如“你们遇到过这种事吗？评论区说说看”。
6.  **文章质量**：生成的内容必须 **完整、通顺、无错别字、标点使用正确、格式规范、逻辑清晰**。绝对避免出现内容空白、缺失、文字不通顺、滥用标点、格式乱码、非通用语言、逻辑混乱等问题。
7.  **格式**：输出纯文本，段落之间空一行。不要用**加粗**、列表、标题符号这些乱七八糟的。

【输出格式】
1.  【诊断报告】：简单说下原文有啥毛病，你打算怎么改。
2.  【优化后标题】：你推荐用的那个标题。
3.  【备选标题】：另外两三个备选。
4.  【优化正文】：改写完的文章，纯文本，段落之间空一行。
5.  【优化说明】：你做了哪些改动，预期效果是啥。

开始整吧。"""


class HeadlineArticleRewriter:
    def __init__(self, root):
        self.root = root
        self.root.title("深度文章优化器 v3.0 - 多版本批量生成版")

        self.api_key = "sk-1f973f482b2241d89c2304a2e4b819db"
        self.api_url = "https://api.deepseek.com/v1/chat/completions"

        self.input_files = []
        self.output_dir = ""          # 用户选择的TXT根目录
        self.docx_output_dir = ""      # 用户选择的DOCX根目录
        self.current_article_data = None
        self.prompts = ArticlePrompts()
        self.api_calls_total = 0
        self.token_usage_total = 0

        self.setup_ui()

    def setup_ui(self):
        main_canvas = tk.Canvas(self.root, borderwidth=0)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=main_canvas.yview)
        main_canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        main_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        main_container = ttk.Frame(main_canvas, padding="15")
        main_canvas.create_window((0, 0), window=main_container, anchor="nw")

        def configure_scroll(event):
            main_canvas.configure(scrollregion=main_canvas.bbox("all"))
        main_container.bind("<Configure>", configure_scroll)

        def on_mousewheel(event):
            main_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        main_canvas.bind_all("<MouseWheel>", on_mousewheel)

        title_label = ttk.Label(main_container, text="📝 深度文章优化器 v3.0",
                                font=("微软雅黑", 16, "bold"))
        title_label.pack(pady=(0, 15))

        # 改写份数设置
        copies_frame = ttk.Frame(main_container)
        copies_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(copies_frame, text="改写份数:").pack(side=tk.LEFT, padx=(0, 5))
        self.copies_var = tk.IntVar(value=3)
        self.copies_spinbox = ttk.Spinbox(copies_frame, from_=1, to=10, textvariable=self.copies_var, width=5)
        self.copies_spinbox.pack(side=tk.LEFT)

        # API统计
        stats_frame = ttk.LabelFrame(main_container, text="📊 API统计", padding="10")
        stats_frame.pack(fill=tk.X, pady=(0, 10))
        self.api_calls_label = ttk.Label(stats_frame, text="API调用: 0次 | 总Token: 0")
        self.api_calls_label.pack()

        file_frame = ttk.LabelFrame(main_container, text="📂 文件管理", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 15))

        file_select_frame = ttk.Frame(file_frame)
        file_select_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(file_select_frame, text="输入文件:").pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(file_select_frame, text="选择TXT文件（可多选）",
                   command=self.select_input_files).pack(side=tk.LEFT, padx=(0, 10))
        self.input_files_label = ttk.Label(file_select_frame, text="未选择文件", foreground="gray")
        self.input_files_label.pack(side=tk.LEFT)

        txt_dir_frame = ttk.Frame(file_frame)
        txt_dir_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(txt_dir_frame, text="TXT输出根目录:").pack(side=tk.LEFT, padx=(0, 5))
        self.output_dir_label = ttk.Label(txt_dir_frame, text="未选择目录", foreground="gray", width=40)
        self.output_dir_label.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(txt_dir_frame, text="选择目录",
                   command=self.select_output_dir).pack(side=tk.LEFT)

        docx_dir_frame = ttk.Frame(file_frame)
        docx_dir_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(docx_dir_frame, text="DOCX输出根目录:").pack(side=tk.LEFT, padx=(0, 5))
        self.docx_output_dir_label = ttk.Label(docx_dir_frame, text="未选择目录（可选）", foreground="gray", width=40)
        self.docx_output_dir_label.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(docx_dir_frame, text="选择目录",
                   command=self.select_docx_output_dir).pack(side=tk.LEFT)

        button_frame = ttk.Frame(file_frame)
        button_frame.pack(fill=tk.X, pady=(0, 10))

        self.generate_btn = ttk.Button(button_frame, text="🚀 开始批量优化",
                                       command=self.start_generation)
        self.generate_btn.pack(side=tk.LEFT, padx=2)

        self.save_btn = ttk.Button(button_frame, text="💾 保存当前文章",
                                   command=self.save_article, state="disabled")
        self.save_btn.pack(side=tk.LEFT, padx=2)

        ttk.Button(button_frame, text="🗑️ 清空预览",
                   command=self.clear_preview).pack(side=tk.LEFT, padx=2)

        self.progress = ttk.Progressbar(file_frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=(5, 0))

        self.status_label = ttk.Label(file_frame, text="就绪", foreground="green")
        self.status_label.pack(pady=(5, 0))

        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        original_tab = ttk.Frame(self.notebook)
        self.notebook.add(original_tab, text="原始内容")
        self.original_text = scrolledtext.ScrolledText(original_tab, wrap=tk.WORD,
                                                       font=("微软雅黑", 10))
        self.original_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        generated_tab = ttk.Frame(self.notebook)
        self.notebook.add(generated_tab, text="优化结果")

        title_frame = ttk.LabelFrame(generated_tab, text="🔥 优化后标题", padding="5")
        title_frame.pack(fill=tk.X, padx=5, pady=5)
        self.title_text = scrolledtext.ScrolledText(title_frame, wrap=tk.WORD,
                                                    font=("微软雅黑", 10), height=4)
        self.title_text.pack(fill=tk.BOTH, expand=True)

        diagnosis_frame = ttk.LabelFrame(generated_tab, text="📋 诊断报告", padding="5")
        diagnosis_frame.pack(fill=tk.X, padx=5, pady=5)
        self.diagnosis_text = scrolledtext.ScrolledText(diagnosis_frame, wrap=tk.WORD,
                                                        font=("微软雅黑", 10), height=6)
        self.diagnosis_text.pack(fill=tk.BOTH, expand=True)

        article_frame = ttk.LabelFrame(generated_tab, text="📖 优化后文章", padding="5")
        article_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))
        self.article_text = scrolledtext.ScrolledText(article_frame, wrap=tk.WORD,
                                                      font=("微软雅黑", 10))
        self.article_text.pack(fill=tk.BOTH, expand=True)

        notes_frame = ttk.LabelFrame(generated_tab, text="💡 优化说明", padding="5")
        notes_frame.pack(fill=tk.X, padx=5, pady=5)
        self.notes_text = scrolledtext.ScrolledText(notes_frame, wrap=tk.WORD,
                                                    font=("微软雅黑", 10), height=4)
        self.notes_text.pack(fill=tk.BOTH, expand=True)

        stats_frame = ttk.LabelFrame(generated_tab, text="📊 统计信息", padding="5")
        stats_frame.pack(fill=tk.X, padx=5, pady=(0, 5))
        self.word_count_label = ttk.Label(stats_frame,
                                          text="原文字数: 0字\n优化后字数: 0字\n原文段落: 0段\n优化后段落: 0段",
                                          font=("微软雅黑", 10))
        self.word_count_label.pack()

    def select_input_files(self):
        files = filedialog.askopenfilenames(
            title="选择要优化的TXT文件（可多选）",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if files:
            self.input_files = list(files)
            file_names = ", ".join([os.path.basename(f) for f in self.input_files[:3]])
            if len(self.input_files) > 3:
                file_names += f" 等{len(self.input_files)}个文件"
            self.input_files_label.config(text=file_names, foreground="green")
            if self.input_files:
                self.load_file_preview(0)

    def select_output_dir(self):
        directory = filedialog.askdirectory(title="选择TXT输出根目录")
        if directory:
            self.output_dir = directory
            self.output_dir_label.config(text=directory, foreground="green")

    def select_docx_output_dir(self):
        directory = filedialog.askdirectory(title="选择DOCX输出根目录")
        if directory:
            self.docx_output_dir = directory
            self.docx_output_dir_label.config(text=directory, foreground="green")

    def load_file_preview(self, file_index):
        if 0 <= file_index < len(self.input_files):
            try:
                with open(self.input_files[file_index], 'r', encoding='utf-8') as f:
                    content = f.read()
                self.original_text.delete(1.0, tk.END)
                self.original_text.insert(1.0, content)
                self.notebook.select(0)
            except Exception as e:
                messagebox.showerror("错误", f"读取文件失败: {str(e)}")

    def extract_main_content(self, file_content):
        lines = file_content.splitlines()
        separator = "============================================================"
        found = False
        main_lines = []
        for line in lines:
            if line.strip() == separator:
                found = True
                continue
            if found:
                main_lines.append(line)
        if found:
            return '\n'.join(main_lines).strip()
        else:
            for i, line in enumerate(lines):
                if line.startswith("内容长度:"):
                    return '\n'.join(lines[i+1:]).strip()
            return file_content.strip()

    def extract_original_title(self, file_content):
        lines = file_content.splitlines()
        for line in lines:
            if line.startswith("标题:"):
                return line[3:].strip()
        return None

    def extract_images(self, file_content):
        lines = file_content.splitlines()
        images = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if re.match(r'^\s*\d+\.\d+:', line):
                if i+1 < len(lines) and lines[i+1].strip().startswith("路径:"):
                    path_line = lines[i+1].strip()
                    match = re.search(r'路径:\s*(.+)', path_line)
                    if match:
                        img_path = match.group(1).strip()
                        images.append(img_path)
                    i += 1
            i += 1
        return images

    def start_generation(self):
        if not self.input_files:
            messagebox.showwarning("警告", "请先选择要优化的文件！")
            return
        if not self.output_dir:
            messagebox.showwarning("警告", "请先选择TXT输出根目录！")
            return

        self.generate_btn.config(state="disabled")
        self.save_btn.config(state="disabled")
        self.progress.start()
        self.status_label.config(text="正在批量优化文章...", foreground="orange")

        thread = threading.Thread(target=self.process_optimization, daemon=True)
        thread.start()

    def process_optimization(self):
        try:
            copies = self.copies_var.get()
            for file_path in self.input_files:
                file_name = os.path.basename(file_path)
                self.status_label.config(text=f"正在优化: {file_name}", foreground="orange")

                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        full_content = f.read()

                    main_content = self.extract_main_content(full_content)
                    image_paths = self.extract_images(full_content)
                    original_title = self.extract_original_title(full_content)

                    if not main_content:
                        self.log_message(f"✗ 无法提取正文，跳过: {file_name}")
                        continue

                    original_paragraphs = self.prompts.count_paragraphs(main_content)

                    # 为当前文件生成一个时间戳，用于文件夹命名
                    folder_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    base_name = os.path.splitext(file_name)[0]

                    # 循环生成多份
                    for copy_idx in range(1, copies + 1):
                        self.log_message(f"正在生成第 {copy_idx}/{copies} 份...")
                        prompt = self.prompts.get_optimization_prompt(main_content, original_paragraphs)
                        result = self.call_deepseek_api(prompt, max_tokens=3000)

                        if result:
                            diagnosis, main_title, alt_titles, raw_article, optimization_notes = self.parse_optimization_result(result)

                            # 清理标题中的数字编号
                            alt_titles = [re.sub(r'^\s*\d+\.\s*', '', t).strip() for t in alt_titles]
                            if main_title:
                                main_title = re.sub(r'^\s*\d+\.\s*', '', main_title).strip()

                            if original_paragraphs < 7:
                                target_para = 7
                            else:
                                target_para = original_paragraphs

                            article = self.clean_article_format(raw_article, target_para)

                            word_count = self.count_chinese_chars(article)
                            original_count = self.count_chinese_chars(main_content)
                            generated_paragraphs = self.prompts.count_paragraphs(article)

                            # 仅预览最新一份（第1份或最后一份，这里预览最后一份）
                            if copy_idx == copies:
                                self.root.after(0, self.update_preview, {
                                    'filename': file_name,
                                    'diagnosis': diagnosis,
                                    'titles': [main_title] + alt_titles,
                                    'article': article,
                                    'optimization_notes': optimization_notes,
                                    'word_count': word_count,
                                    'original_word_count': original_count,
                                    'original_paragraphs': original_paragraphs,
                                    'generated_paragraphs': generated_paragraphs
                                })

                            # 保存TXT
                            txt_path = self.save_article_to_file(
                                file_name, full_content, main_content,
                                [main_title] + alt_titles, article,
                                original_count, word_count,
                                diagnosis=diagnosis,
                                optimization_notes=optimization_notes,
                                folder_timestamp=folder_timestamp,
                                copy_idx=copy_idx
                            )
                            self.log_message(f"✓ 已保存TXT: {txt_path}")

                            # 保存DOCX
                            if self.docx_output_dir and DOCX_AVAILABLE:
                                docx_path = self.generate_docx(
                                    file_name, article, image_paths,
                                    [main_title] + alt_titles, generated_paragraphs,
                                    original_title,
                                    folder_timestamp=folder_timestamp,
                                    copy_idx=copy_idx
                                )
                                if docx_path:
                                    self.log_message(f"✓ 已保存DOCX: {docx_path}")
                                else:
                                    self.log_message(f"✗ DOCX生成失败: {file_name} 第{copy_idx}份")
                            elif not DOCX_AVAILABLE:
                                self.log_message("⚠️ python-docx未安装，跳过DOCX生成")

                            # 每次调用后稍作延时
                            time.sleep(2)
                        else:
                            self.log_message(f"✗ 第{copy_idx}份优化失败: {file_name}")

                    self.log_message(f"✓ 完成文件: {file_name}，共生成{copies}份")

                except Exception as e:
                    self.log_message(f"处理失败 {file_name}: {str(e)}")

            self.root.after(0, self.generation_complete)

        except Exception as e:
            self.root.after(0, self.generation_error, str(e))

    def parse_optimization_result(self, result):
        diagnosis = ""
        main_title = ""
        alt_titles = []
        article = ""
        optimization_notes = ""

        if "【诊断报告】" in result:
            diagnosis_section = result.split("【诊断报告】")[1].split("【优化后标题】")[0].strip()
            diagnosis = diagnosis_section

        if "【优化后标题】" in result:
            title_section = result.split("【优化后标题】")[1].split("【备选标题】")[0].strip()
            main_title = title_section

        if "【备选标题】" in result:
            alt_section = result.split("【备选标题】")[1].split("【优化正文】")[0].strip()
            alt_titles = [line.strip() for line in alt_section.split('\n') if line.strip()]

        if "【优化正文】" in result:
            article_section = result.split("【优化正文】")[1]
            if "【优化说明】" in article_section:
                article = article_section.split("【优化说明】")[0].strip()
                notes_section = article_section.split("【优化说明】")[1].strip()
                optimization_notes = notes_section
            else:
                article = article_section.strip()

        return diagnosis, main_title, alt_titles, article, optimization_notes

    def clean_article_format(self, article, target_para=None):
        article = re.sub(r'\*\*(.+?)\*\*', r'\1', article)
        article = re.sub(r'\*(.+?)\*', r'\1', article)
        article = re.sub(r'#+\s*', '', article)
        article = re.sub(r'`(.+?)`', r'\1', article)
        article = re.sub(r'^\s*[0-9]+\.\s*', '', article, flags=re.MULTILINE)
        article = re.sub(r'^\s*[-*•]\s*', '', article, flags=re.MULTILINE)

        if target_para is not None:
            lines = [line.strip() for line in article.splitlines() if line.strip()]
            if not lines:
                return ""
            if len(lines) <= target_para:
                new_paras = lines
            else:
                lines_per_para = math.ceil(len(lines) / target_para)
                new_paras = []
                for i in range(0, len(lines), lines_per_para):
                    para = ' '.join(lines[i:i+lines_per_para])
                    new_paras.append(para)
            return '\n\n'.join(new_paras)
        else:
            raw_paras = [p.strip() for p in re.split(r'\n\s*\n', article.strip()) if p.strip()]
            cleaned_paras = []
            for para in raw_paras:
                para = re.sub(r'\s+', ' ', para.replace('\n', ' ')).strip()
                if para:
                    cleaned_paras.append(para)
            return '\n\n'.join(cleaned_paras)

    def generate_docx(self, filename, article, image_paths, titles, num_paragraphs, original_title,
                      folder_timestamp, copy_idx):
        """
        生成DOCX文件，保存在 docx根目录/原文件名_时间戳/ 下，文件名包含份数序号
        """
        if not self.docx_output_dir:
            return None

        base_name = os.path.splitext(filename)[0]
        folder_name = f"{base_name}_{folder_timestamp}"
        docx_dir = os.path.join(self.docx_output_dir, "docx", folder_name)
        os.makedirs(docx_dir, exist_ok=True)

        doc = Document()

        def set_font(paragraph, bold=False, size=12):
            for run in paragraph.runs:
                run.font.name = '宋体'
                run.font.size = Pt(size)
                run.font.bold = bold

        if original_title:
            p = doc.add_paragraph(original_title)
            set_font(p, bold=True, size=14)
            doc.add_paragraph()

        if titles:
            p = doc.add_paragraph("推荐标题：")
            set_font(p, bold=False)
            for title in titles:
                p = doc.add_paragraph(title)
                set_font(p, bold=True)
            doc.add_paragraph()

        paragraphs = article.split('\n\n')
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        images_to_use = image_paths[:7]
        num_images = len(images_to_use)
        if num_images > 0 and num_paragraphs > 0:
            insert_positions = []
            for i in range(1, num_images + 1):
                pos = math.floor(i * (num_paragraphs + 1) / (num_images + 1))
                insert_positions.append(pos)
            insert_positions = [min(p, num_paragraphs) for p in insert_positions]
            insert_positions = sorted(set(insert_positions))
            if len(insert_positions) < num_images:
                while len(insert_positions) < num_images:
                    insert_positions.append(num_paragraphs)
                insert_positions = sorted(set(insert_positions))[:num_images]
        else:
            insert_positions = []

        img_index = 0
        for idx, para_text in enumerate(paragraphs, start=1):
            p = doc.add_paragraph(para_text)
            set_font(p, bold=False)

            if img_index < len(insert_positions) and idx == insert_positions[img_index]:
                img_path = images_to_use[img_index]
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.0))
                        doc.add_paragraph()
                    except Exception as e:
                        self.log_message(f"插入图片失败 {img_path}: {e}")
                else:
                    self.log_message(f"图片文件不存在: {img_path}")
                img_index += 1

        output_filename = f"优化_{base_name}_{copy_idx}_{folder_timestamp}.docx"
        output_path = os.path.join(docx_dir, output_filename)

        try:
            doc.save(output_path)
            return output_path
        except Exception as e:
            self.log_message(f"保存DOCX失败: {e}")
            return None

    def update_preview(self, data):
        self.title_text.delete(1.0, tk.END)
        self.diagnosis_text.delete(1.0, tk.END)
        self.article_text.delete(1.0, tk.END)
        self.notes_text.delete(1.0, tk.END)

        if data['titles']:
            titles_text = '\n'.join(data['titles'][:5])
            self.title_text.insert(1.0, titles_text)

        if data.get('diagnosis'):
            self.diagnosis_text.insert(1.0, data['diagnosis'])

        self.article_text.insert(1.0, data['article'])

        if data.get('optimization_notes'):
            self.notes_text.insert(1.0, data['optimization_notes'])

        word_count = data['word_count']
        original_count = data.get('original_word_count', 0)
        diff = word_count - original_count if original_count > 0 else 0
        diff_text = f" (+{diff}字)" if diff > 0 else f" ({diff}字)" if diff < 0 else ""

        self.word_count_label.config(
            text=f"原文字数: {original_count}字\n优化后字数: {word_count}字{diff_text}\n"
                 f"原文段落: {data.get('original_paragraphs', 0)}段\n优化后段落: {data.get('generated_paragraphs', 0)}段"
        )

        self.notebook.select(1)
        self.current_article_data = data
        self.save_btn.config(state="normal")
        self.progress.stop()
        self.status_label.config(text="优化完成！", foreground="green")

    def save_article_to_file(self, filename, full_content, main_content, titles, article,
                             original_count, word_count, diagnosis="", optimization_notes="",
                             folder_timestamp=None, copy_idx=1):
        """
        保存TXT文件到 txt根目录/原文件名_时间戳/ 下，文件名包含份数序号
        """
        if not self.output_dir:
            return None

        base_name = os.path.splitext(filename)[0]
        if folder_timestamp is None:
            folder_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        folder_name = f"{base_name}_{folder_timestamp}"
        txt_dir = os.path.join(self.output_dir, "txt", folder_name)
        os.makedirs(txt_dir, exist_ok=True)

        output_filename = f"优化_{base_name}_{copy_idx}_{folder_timestamp}.txt"
        output_path = os.path.join(txt_dir, output_filename)

        separator = "============================================================"
        if separator in full_content:
            header = full_content.split(separator)[0].strip()
            output_content = f"{header}\n{separator}\n\n【优化后文章】\n\n"
            output_content += f"优化后标题：{titles[0] if titles else '无'}\n"
            if len(titles) > 1:
                output_content += f"备选标题：{', '.join(titles[1:])}\n"
            output_content += f"\n{article}\n"
            if diagnosis:
                output_content += f"\n【诊断报告】\n{diagnosis}\n"
            if optimization_notes:
                output_content += f"\n【优化说明】\n{optimization_notes}\n"
        else:
            output_content = f"""生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
原始文件：{filename}
原文字数：{original_count}字
优化后字数：{word_count}字
{'=' * 60}

🔥 优化后标题：
{chr(10).join(titles[:5])}

{'=' * 60}
"""
            if diagnosis:
                output_content += f"""
【诊断报告】
{diagnosis}

{'=' * 60}
"""
            output_content += f"""
📖 优化后文章：

{article}
"""
            if optimization_notes:
                output_content += f"""

【优化说明】
{optimization_notes}
"""
            output_content += f"""
{'=' * 60}

📋 原始内容摘要：
{main_content[:500]}...
"""

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(output_content)
        return output_path

    def save_article(self):
        if not self.output_dir:
            messagebox.showwarning("警告", "请先选择TXT输出根目录！")
            return
        if not self.current_article_data:
            messagebox.showwarning("警告", "没有可保存的内容！")
            return

        titles = [line.strip() for line in self.title_text.get(1.0, tk.END).strip().split('\n') if line.strip()]
        article = self.article_text.get(1.0, tk.END).strip()
        word_count = self.count_chinese_chars(article)

        try:
            with open(self.current_article_data['filename_full'], 'r', encoding='utf-8') as f:
                full_content = f.read()
        except:
            full_content = ""

        output_path = self.save_article_to_file(
            self.current_article_data['filename'],
            full_content,
            self.current_article_data.get('original', ''),
            titles,
            article,
            self.current_article_data.get('original_word_count', 0),
            word_count,
            self.current_article_data.get('diagnosis', ''),
            self.current_article_data.get('optimization_notes', '')
        )
        messagebox.showinfo("成功", f"文章已保存到：\n{output_path}")

    def generation_complete(self):
        self.progress.stop()
        self.generate_btn.config(state="normal")
        self.status_label.config(text="✅ 批量优化完成！", foreground="green")
        self.api_calls_label.config(
            text=f"API调用: {self.api_calls_total}次 | 总Token: {self.token_usage_total}"
        )
        messagebox.showinfo("完成",
                            f"批量优化完成！\n处理文件: {len(self.input_files)}个\nAPI调用: {self.api_calls_total}次")

    def generation_error(self, error_msg):
        self.progress.stop()
        self.generate_btn.config(state="normal")
        self.status_label.config(text=f"❌ 处理失败: {error_msg}", foreground="red")
        messagebox.showerror("错误", f"处理过程中出现错误:\n{error_msg}")

    def count_chinese_chars(self, text):
        chinese_chars = re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', text)
        return len(chinese_chars)

    def call_deepseek_api(self, prompt, max_tokens=2000):
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_key}'
        }
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system",
                 "content": "你是一个有十年经验的爆文编辑，擅长把普通文章改成高传播、高互动的爆款。你说话直接、接地气，文章风格口语化，讨厌AI腔。你会使用大白话、短句，避免使用‘与此同时’之类的连接词。你会刻意制造争议点，吸引读者评论。你生成的文章必须内容完整、通顺、无错别字、标点使用正确、格式规范、逻辑清晰，避免任何影响阅读体验的质量问题。输出必须纯文本，段落之间空一行，段落数量必须严格控制在用户要求的范围内。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.85,
            "max_tokens": max_tokens
        }

        try:
            response = requests.post(self.api_url, headers=headers, json=payload, timeout=120)
            response.raise_for_status()
            result = response.json()

            if 'choices' in result and len(result['choices']) > 0:
                content = result['choices'][0]['message']['content'].strip()
                self.api_calls_total += 1
                estimated_tokens = len(prompt) // 3 + max_tokens // 2
                self.token_usage_total += estimated_tokens
                self.log_message(f"API调用成功，返回长度：{len(content)}字符")
                return content
            else:
                self.log_message(f"API响应异常")
                return None

        except Exception as e:
            self.log_message(f"API调用失败: {e}")
            return None

    def log_message(self, message):
        timestamp = datetime.now().strftime('%H:%M:%S')
        print(f"[{timestamp}] {message}")

    def clear_preview(self):
        self.original_text.delete(1.0, tk.END)
        self.title_text.delete(1.0, tk.END)
        self.diagnosis_text.delete(1.0, tk.END)
        self.article_text.delete(1.0, tk.END)
        self.notes_text.delete(1.0, tk.END)
        self.word_count_label.config(text="原文字数: 0字\n优化后字数: 0字\n原文段落: 0段\n优化后段落: 0段")
        self.save_btn.config(state="disabled")
        self.current_article_data = None
        self.status_label.config(text="就绪", foreground="green")
        self.notebook.select(0)


def main():
    root = tk.Tk()
    root.geometry("1000x800")
    app = HeadlineArticleRewriter(root)
    try:
        root.iconbitmap(default='icon.ico')
    except:
        pass
    root.mainloop()


if __name__ == "__main__":
    main()